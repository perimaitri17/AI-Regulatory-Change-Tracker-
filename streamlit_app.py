import streamlit as st
import pandas as pd
import requests
from datetime import datetime, timedelta
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import json
import hashlib
from dataclasses import dataclass
from typing import List, Dict, Optional
import re
from io import BytesIO
import docx
from docx.shared import RGBColor
from docx import Document
import time

# Configure page
st.set_page_config(
    page_title="Pharma Regulatory Intelligence Platform",
    page_icon="🏥",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin: 0.5rem 0;
    }
    .alert-high {
        background-color: #ff4444;
        padding: 1rem;
        border-radius: 8px;
        color: white;
        margin: 0.5rem 0;
    }
    .alert-medium {
        background-color: #ffaa00;
        padding: 1rem;
        border-radius: 8px;
        color: white;
        margin: 0.5rem 0;
    }
    .alert-low {
        background-color: #00aa00;
        padding: 1rem;
        border-radius: 8px;
        color: white;
        margin: 0.5rem 0;
    }
    .sidebar .sidebar-content {
        background: linear-gradient(180deg, #f8f9fa 0%, #e9ecef 100%);
    }
</style>
""", unsafe_allow_html=True)

@dataclass
class RegulatorySource:
    name: str
    url: str
    region: str
    categories: List[str]
    update_frequency: str
    api_endpoint: Optional[str] = None

@dataclass
class ProductMapping:
    product_name: str
    therapeutic_area: str
    formulation: str
    indications: List[str]
    regulatory_pathways: List[str]
    regions: List[str]

@dataclass
class RegulatoryChange:
    id: str
    title: str
    description: str
    source: str
    region: str
    category: str
    therapeutic_area: str
    impact_level: str
    effective_date: datetime
    products_affected: List[str]
    action_required: str
    url: str
    detected_date: datetime

class RegulatoryIntelligencePlatform:
    def __init__(self):
        self.regulatory_sources = self._initialize_sources()
        self.product_mappings = self._initialize_product_mappings()
        self.mock_changes = self._generate_mock_changes()
        
    def _initialize_sources(self) -> List[RegulatorySource]:
        return [
            # EU Sources
            RegulatorySource(
                name="European Medicines Agency (EMA)",
                url="https://www.ema.europa.eu",
                region="EU",
                categories=["Drug Approval", "Safety Updates", "Guidelines"],
                update_frequency="Daily"
            ),
            RegulatorySource(
                name="European Commission - Health",
                url="https://ec.europa.eu/health",
                region="EU",
                categories=["Medical Devices", "Pharmaceuticals", "Clinical Trials"],
                update_frequency="Weekly"
            ),
            RegulatorySource(
                name="EudraLex",
                url="https://ec.europa.eu/health/documents/eudralex",
                region="EU",
                categories=["GMP", "GDP", "Regulatory Guidelines"],
                update_frequency="Monthly"
            ),
            
            # US Sources
            RegulatorySource(
                name="FDA - Drugs",
                url="https://www.fda.gov/drugs",
                region="US",
                categories=["Drug Approval", "Safety Alerts", "Guidance Documents"],
                update_frequency="Daily"
            ),
            RegulatorySource(
                name="FDA - CDER",
                url="https://www.fda.gov/about-fda/center-drug-evaluation-and-research-cder",
                region="US",
                categories=["Clinical Trials", "Drug Development", "Regulatory Science"],
                update_frequency="Weekly"
            ),
            RegulatorySource(
                name="Federal Register",
                url="https://www.federalregister.gov",
                region="US",
                categories=["Final Rules", "Proposed Rules", "Notices"],
                update_frequency="Daily"
            ),
            
            # India Sources
            RegulatorySource(
                name="Central Drugs Standard Control Organization (CDSCO)",
                url="https://cdsco.gov.in",
                region="India",
                categories=["Drug Approval", "Clinical Trials", "Import Licenses"],
                update_frequency="Weekly"
            ),
            RegulatorySource(
                name="Ministry of Health - India",
                url="https://mohfw.gov.in",
                region="India",
                categories=["Health Policies", "Drug Policies", "Medical Devices"],
                update_frequency="Monthly"
            ),
            RegulatorySource(
                name="Indian Pharmacopoeia Commission",
                url="https://ipc.gov.in",
                region="India",
                categories=["Standards", "Quality Control", "Pharmacovigilance"],
                update_frequency="Quarterly"
            ),
            
            # International Sources
            RegulatorySource(
                name="WHO",
                url="https://www.who.int",
                region="International",
                categories=["Global Health", "Drug Safety", "Standards"],
                update_frequency="Weekly"
            ),
            RegulatorySource(
                name="ICH Guidelines",
                url="https://www.ich.org",
                region="International",
                categories=["Harmonized Guidelines", "Quality", "Safety", "Efficacy"],
                update_frequency="Quarterly"
            )
        ]
    
    def _initialize_product_mappings(self) -> List[ProductMapping]:
        return [
            ProductMapping(
                product_name="Atorvastatin Tablets",
                therapeutic_area="Cardiovascular",
                formulation="Oral Solid Dosage",
                indications=["Hypercholesterolemia", "Cardiovascular Risk Reduction"],
                regulatory_pathways=["ANDA", "Generic MA", "CDSCO Approval"],
                regions=["US", "EU", "India"]
            ),
            ProductMapping(
                product_name="Insulin Glargine Injection",
                therapeutic_area="Diabetes",
                formulation="Injectable",
                indications=["Type 1 Diabetes", "Type 2 Diabetes"],
                regulatory_pathways=["BLA", "Biosimilar MA", "New Drug Approval"],
                regions=["US", "EU", "India"]
            ),
            ProductMapping(
                product_name="Montelukast Oral Granules",
                therapeutic_area="Respiratory",
                formulation="Oral Granules",
                indications=["Asthma", "Allergic Rhinitis"],
                regulatory_pathways=["ANDA", "Generic MA", "Fixed Dose Combination"],
                regions=["US", "EU", "India"]
            ),
            ProductMapping(
                product_name="Adalimumab Biosimilar",
                therapeutic_area="Immunology",
                formulation="Pre-filled Syringe",
                indications=["Rheumatoid Arthritis", "Inflammatory Bowel Disease"],
                regulatory_pathways=["BLA", "Biosimilar MA", "Biological Product"],
                regions=["US", "EU", "India"]
            ),
            ProductMapping(
                product_name="Paclitaxel Injection",
                therapeutic_area="Oncology",
                formulation="Injectable",
                indications=["Breast Cancer", "Lung Cancer", "Ovarian Cancer"],
                regulatory_pathways=["ANDA", "Generic MA", "Cytotoxic Drug Approval"],
                regions=["US", "EU", "India"]
            )
        ]
    
    def _generate_mock_changes(self) -> List[RegulatoryChange]:
        changes = []
        base_date = datetime.now() - timedelta(days=30)
        
        mock_data = [
            {
                "title": "Updated FDA Guidance on Biosimilar Development for Monoclonal Antibodies",
                "description": "FDA has released revised guidance on analytical and clinical considerations for biosimilar monoclonal antibodies, including new requirements for immunogenicity studies.",
                "source": "FDA - CDER",
                "region": "US",
                "category": "Guidance Documents",
                "therapeutic_area": "Immunology",
                "impact_level": "High",
                "products_affected": ["Adalimumab Biosimilar"],
                "action_required": "Review immunogenicity study protocols, update regulatory strategy documents",
                "url": "https://www.fda.gov/regulatory-information/search-fda-guidance-documents"
            },
            {
                "title": "EMA New Quality Requirements for Oral Solid Dosage Forms",
                "description": "European Medicines Agency introduces enhanced quality standards for immediate-release tablets and capsules, focusing on dissolution testing and stability studies.",
                "source": "European Medicines Agency (EMA)",
                "region": "EU",
                "category": "Guidelines",
                "therapeutic_area": "Cardiovascular",
                "impact_level": "Medium",
                "products_affected": ["Atorvastatin Tablets"],
                "action_required": "Update dissolution methods, revise stability protocols",
                "url": "https://www.ema.europa.eu/en/human-regulatory/research-development"
            },
            {
                "title": "CDSCO Revised Schedule Y for Clinical Trials",
                "description": "New clinical trial requirements for diabetes medications in Indian population, including specific bioequivalence criteria for insulin formulations.",
                "source": "CDSCO",
                "region": "India",
                "category": "Clinical Trials",
                "therapeutic_area": "Diabetes",
                "impact_level": "High",
                "products_affected": ["Insulin Glargine Injection"],
                "action_required": "Redesign bioequivalence studies, update clinical protocols",
                "url": "https://cdsco.gov.in/opencms/opencms/en/Clinical-Trial/"
            },
            {
                "title": "WHO Updated Guidelines on Oncology Drug Development",
                "description": "World Health Organization releases updated guidance on development and approval of generic oncology drugs with enhanced bioequivalence requirements.",
                "source": "WHO",
                "region": "International",
                "category": "Global Health",
                "therapeutic_area": "Oncology",
                "impact_level": "Medium",
                "products_affected": ["Paclitaxel Injection"],
                "action_required": "Review bioequivalence study design, assess global impact",
                "url": "https://www.who.int/medicines/areas/quality_safety"
            },
            {
                "title": "ICH M10 Guideline on Bioanalytical Method Validation - Final",
                "description": "Harmonized guideline on bioanalytical method validation replacing regional guidelines, affecting all therapeutic areas.",
                "source": "ICH Guidelines",
                "region": "International",
                "category": "Harmonized Guidelines",
                "therapeutic_area": "All",
                "impact_level": "High",
                "products_affected": ["All Products"],
                "action_required": "Update all bioanalytical validation procedures, retrain analytical teams",
                "url": "https://www.ich.org/page/multidisciplinary-guidelines"
            }
        ]
        
        for i, data in enumerate(mock_data):
            change_id = hashlib.md5(f"{data['title']}{i}".encode()).hexdigest()[:8]
            changes.append(RegulatoryChange(
                id=change_id,
                title=data["title"],
                description=data["description"],
                source=data["source"],
                region=data["region"],
                category=data["category"],
                therapeutic_area=data["therapeutic_area"],
                impact_level=data["impact_level"],
                effective_date=base_date + timedelta(days=i*5 + 30),
                products_affected=data["products_affected"],
                action_required=data["action_required"],
                url=data["url"],
                detected_date=base_date + timedelta(days=i*3)
            ))
        
        return changes

def create_document_comparison():
    """Create a sample document comparison showing before/after regulatory updates"""
    
    # Create original document
    original_doc = Document()
    original_doc.add_heading('Bioanalytical Method Validation Protocol', 0)
    original_doc.add_heading('1. Introduction', level=1)
    p1 = original_doc.add_paragraph(
        "This protocol describes the validation of bioanalytical methods according to "
        "FDA Guidance for Industry: Bioanalytical Method Validation (2018) and "
        "EMA Guideline on bioanalytical method validation (2011)."
    )
    
    original_doc.add_heading('2. Validation Parameters', level=1)
    original_doc.add_paragraph("• Selectivity and Specificity")
    original_doc.add_paragraph("• Accuracy and Precision")
    original_doc.add_paragraph("• Calibration Curve and Range")
    original_doc.add_paragraph("• Lower Limit of Quantification (LLOQ)")
    original_doc.add_paragraph("• Stability")
    
    # Create updated document with tracked changes
    updated_doc = Document()
    updated_doc.add_heading('Bioanalytical Method Validation Protocol', 0)
    updated_doc.add_heading('1. Introduction', level=1)
    
    # Updated paragraph with tracked changes indication
    p1_updated = updated_doc.add_paragraph()
    p1_updated.add_run("This protocol describes the validation of bioanalytical methods according to ")
    
    # Simulate tracked changes
    run_deleted = p1_updated.add_run("FDA Guidance for Industry: Bioanalytical Method Validation (2018) and EMA Guideline on bioanalytical method validation (2011)")
    run_deleted.font.color.rgb = RGBColor(255, 0, 0)  # Red for deleted text
    
    run_added = p1_updated.add_run("ICH M10 Guideline on Bioanalytical Method Validation (2022) which harmonizes global requirements")
    run_added.font.color.rgb = RGBColor(0, 128, 0)  # Green for added text
    p1_updated.add_run(".")
    
    updated_doc.add_heading('2. Validation Parameters', level=1)
    updated_doc.add_paragraph("• Selectivity and Specificity")
    updated_doc.add_paragraph("• Accuracy and Precision")
    updated_doc.add_paragraph("• Calibration Curve and Range")
    updated_doc.add_paragraph("• Lower Limit of Quantification (LLOQ)")
    
    # Add new requirement
    new_req = updated_doc.add_paragraph()
    new_req.add_run("• ").font.color.rgb = RGBColor(0, 0, 0)
    new_req.add_run("Matrix Effect Assessment (New ICH M10 requirement)").font.color.rgb = RGBColor(0, 128, 0)
    
    updated_doc.add_paragraph("• Stability")
    
    return original_doc, updated_doc

def main():
    st.markdown('<h1 class="main-header">🏥 Pharma Regulatory Intelligence Platform</h1>', unsafe_allow_html=True)
    
    # Initialize platform
    if 'platform' not in st.session_state:
        st.session_state.platform = RegulatoryIntelligencePlatform()
    
    platform = st.session_state.platform
    
    # Sidebar for filters and configuration
    with st.sidebar:
        st.header("🎛️ Configuration")
        
        # Region filter
        selected_regions = st.multiselect(
            "Select Regions",
            ["US", "EU", "India", "International"],
            default=["US", "EU", "India"]
        )
        
        # Therapeutic area filter
        therapeutic_areas = ["All", "Cardiovascular", "Diabetes", "Respiratory", "Immunology", "Oncology"]
        selected_therapeutic_areas = st.multiselect(
            "Therapeutic Areas",
            therapeutic_areas,
            default=["All"]
        )
        
        # Impact level filter
        impact_levels = st.multiselect(
            "Impact Level",
            ["High", "Medium", "Low"],
            default=["High", "Medium"]
        )
        
        st.header("🔄 Automation Settings")
        
        # Workflow automation toggles
        auto_document_update = st.checkbox("Auto Document Updates", value=True)
        auto_notifications = st.checkbox("Real-time Notifications", value=True)
        auto_compliance_check = st.checkbox("Compliance Monitoring", value=True)
        
        # Integration options
        st.subheader("📊 Integrations")
        st.selectbox("Document Management", ["SharePoint", "Google Drive", "OneDrive", "Custom"])
        st.selectbox("Notification Channel", ["Email", "Slack", "Teams", "Webhook"])
        st.selectbox("Project Management", ["Jira", "Azure DevOps", "Monday.com", "Custom"])
    
    # Main dashboard tabs
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "📊 Dashboard", 
        "🔍 Change Tracking", 
        "📋 Product Impact", 
        "📄 Document Updates", 
        "⚡ Workflow Automation",
        "🎯 Compliance Monitor"
    ])
    
    with tab1:
        st.header("Regulatory Intelligence Dashboard")
        
        # Key metrics
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.markdown("""
            <div class="metric-card">
                <h3>124</h3>
                <p>Active Changes</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            st.markdown("""
            <div class="metric-card">
                <h3>18</h3>
                <p>High Impact</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col3:
            st.markdown("""
            <div class="metric-card">
                <h3>67</h3>
                <p>Products Affected</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col4:
            st.markdown("""
            <div class="metric-card">
                <h3>12</h3>
                <p>Actions Required</p>
            </div>
            """, unsafe_allow_html=True)
        
        # Recent high-impact changes
        st.subheader("🚨 High-Impact Changes")
        
        high_impact_changes = [change for change in platform.mock_changes if change.impact_level == "High"]
        
        for change in high_impact_changes[:3]:
            alert_class = f"alert-{change.impact_level.lower()}"
            st.markdown(f"""
            <div class="{alert_class}">
                <h4>{change.title}</h4>
                <p><strong>Region:</strong> {change.region} | <strong>Therapeutic Area:</strong> {change.therapeutic_area}</p>
                <p><strong>Products Affected:</strong> {', '.join(change.products_affected)}</p>
                <p><strong>Action Required:</strong> {change.action_required}</p>
            </div>
            """, unsafe_allow_html=True)
        
        # Regional distribution chart
        col1, col2 = st.columns(2)
        
        with col1:
            region_data = {}
            for change in platform.mock_changes:
                region_data[change.region] = region_data.get(change.region, 0) + 1
            
            fig_region = px.pie(
                values=list(region_data.values()),
                names=list(region_data.keys()),
                title="Changes by Region",
                color_discrete_sequence=px.colors.qualitative.Set3
            )
            st.plotly_chart(fig_region, use_container_width=True)
        
        with col2:
            therapeutic_data = {}
            for change in platform.mock_changes:
                therapeutic_data[change.therapeutic_area] = therapeutic_data.get(change.therapeutic_area, 0) + 1
            
            fig_therapeutic = px.bar(
                x=list(therapeutic_data.keys()),
                y=list(therapeutic_data.values()),
                title="Changes by Therapeutic Area",
                color=list(therapeutic_data.values()),
                color_continuous_scale="viridis"
            )
            st.plotly_chart(fig_therapeutic, use_container_width=True)
    
    with tab2:
        st.header("🔍 Regulatory Change Tracking")
        
        # Filter changes based on sidebar selections
        filtered_changes = platform.mock_changes
        if "All" not in selected_therapeutic_areas:
            filtered_changes = [c for c in filtered_changes if c.therapeutic_area in selected_therapeutic_areas]
        if selected_regions:
            filtered_changes = [c for c in filtered_changes if c.region in selected_regions]
        if impact_levels:
            filtered_changes = [c for c in filtered_changes if c.impact_level in impact_levels]
        
        # Create DataFrame for display
        changes_data = []
        for change in filtered_changes:
            changes_data.append({
                "ID": change.id,
                "Title": change.title,
                "Region": change.region,
                "Category": change.category,
                "Therapeutic Area": change.therapeutic_area,
                "Impact Level": change.impact_level,
                "Effective Date": change.effective_date.strftime("%Y-%m-%d"),
                "Products Affected": len(change.products_affected),
                "Action Required": "Yes" if change.action_required else "No"
            })
        
        df_changes = pd.DataFrame(changes_data)
        
        if not df_changes.empty:
            st.dataframe(
                df_changes,
                use_container_width=True,
                hide_index=True
            )
            
            # Detailed view
            if st.button("View Detailed Analysis"):
                for change in filtered_changes[:3]:  # Show first 3 for demo
                    with st.expander(f"📋 {change.title}"):
                        col1, col2 = st.columns(2)
                        with col1:
                            st.write(f"**Source:** {change.source}")
                            st.write(f"**Region:** {change.region}")
                            st.write(f"**Category:** {change.category}")
                            st.write(f"**Impact Level:** {change.impact_level}")
                        with col2:
                            st.write(f"**Therapeutic Area:** {change.therapeutic_area}")
                            st.write(f"**Effective Date:** {change.effective_date.strftime('%Y-%m-%d')}")
                            st.write(f"**Detected:** {change.detected_date.strftime('%Y-%m-%d')}")
                        
                        st.write(f"**Description:** {change.description}")
                        st.write(f"**Products Affected:** {', '.join(change.products_affected)}")
                        st.write(f"**Action Required:** {change.action_required}")
                        
                        if st.button(f"Generate Action Plan for {change.id}"):
                            st.success(f"✅ Action plan generated for {change.title}")
        else:
            st.info("No changes match the current filters.")
    
    with tab3:
        st.header("📋 Product Impact Analysis")
        
        # Product mapping table
        st.subheader("Product Portfolio Mapping")
        
        product_data = []
        for product in platform.product_mappings:
            # Count affected changes
            affected_changes = 0
            for change in platform.mock_changes:
                if (product.product_name in change.products_affected or 
                    change.therapeutic_area == product.therapeutic_area or 
                    "All Products" in change.products_affected):
                    affected_changes += 1
            
            product_data.append({
                "Product": product.product_name,
                "Therapeutic Area": product.therapeutic_area,
                "Formulation": product.formulation,
                "Regions": ", ".join(product.regions),
                "Regulatory Pathways": ", ".join(product.regulatory_pathways),
                "Active Changes": affected_changes,
                "Risk Level": "High" if affected_changes > 2 else "Medium" if affected_changes > 0 else "Low"
            })
        
        df_products = pd.DataFrame(product_data)
        st.dataframe(df_products, use_container_width=True, hide_index=True)
        
        # Product-specific analysis
        st.subheader("Product-Specific Impact Analysis")
        selected_product = st.selectbox("Select Product for Detailed Analysis", 
                                      [p.product_name for p in platform.product_mappings])
        
        if selected_product:
            product = next(p for p in platform.product_mappings if p.product_name == selected_product)
            
            col1, col2 = st.columns(2)
            with col1:
                st.write(f"**Therapeutic Area:** {product.therapeutic_area}")
                st.write(f"**Formulation:** {product.formulation}")
                st.write(f"**Indications:** {', '.join(product.indications)}")
            with col2:
                st.write(f"**Regulatory Pathways:** {', '.join(product.regulatory_pathways)}")
                st.write(f"**Markets:** {', '.join(product.regions)}")
            
            # Show relevant changes
            relevant_changes = [
                change for change in platform.mock_changes
                if (selected_product in change.products_affected or 
                    change.therapeutic_area == product.therapeutic_area or
                    "All Products" in change.products_affected)
            ]
            
            if relevant_changes:
                st.write(f"**Regulatory Changes Affecting {selected_product}:**")
                for change in relevant_changes:
                    impact_color = {"High": "🔴", "Medium": "🟡", "Low": "🟢"}[change.impact_level]
                    st.write(f"{impact_color} {change.title} ({change.region})")
    
    with tab4:
        st.header("📄 Automated Document Updates")
        
        st.subheader("Document Change Analysis")
        st.write("""
        This section demonstrates how regulatory changes automatically trigger document updates 
        with tracked changes for review and approval.
        """)
        
        # Simulate document comparison
        if st.button("Generate Document Comparison"):
            with st.spinner("Analyzing regulatory changes and updating documents..."):
                time.sleep(2)  # Simulate processing
                
                st.success("✅ Document analysis complete!")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.subheader("📄 Original Document")
                    st.code("""
Bioanalytical Method Validation Protocol

1. Introduction
This protocol describes the validation of bioanalytical 
methods according to FDA Guidance for Industry: 
Bioanalytical Method Validation (2018) and EMA 
Guideline on bioanalytical method validation (2011).

2. Validation Parameters
• Selectivity and Specificity
• Accuracy and Precision  
• Calibration Curve and Range
• Lower Limit of Quantification (LLOQ)
• Stability
                    """, language="markdown")
                
                with col2:
                    st.subheader("📝 Updated Document (Tracked Changes)")
                    st.markdown("""
**Bioanalytical Method Validation Protocol**

**1. Introduction**  
This protocol describes the validation of bioanalytical methods according to ~~FDA Guidance for Industry: Bioanalytical Method Validation (2018) and EMA Guideline on bioanalytical method validation (2011)~~ **ICH M10 Guideline on Bioanalytical Method Validation (2022) which harmonizes global requirements**.

**2. Validation Parameters**
• Selectivity and Specificity  
• Accuracy and Precision  
• Calibration Curve and Range  
• Lower Limit of Quantification (LLOQ)  
• **Matrix Effect Assessment (New ICH M10 requirement)**  
• Stability

**Key Changes:**
- ✅ Updated reference to new ICH M10 guideline
- ✅ Added matrix effect assessment requirement
- ✅ Harmonized global requirements integrated
                    """)
        
        # Document workflow status
        st.subheader("📊 Document Update Status")
        
        doc_status_data = [
            {"Document": "Bioanalytical Method Validation Protocol", "Status": "Updated", "Reviewer": "John Smith", "Due Date": "2024-09-01"},
            {"Document": "Clinical Study Protocol Template", "Status": "In Review", "Reviewer": "Sarah Johnson", "Due Date": "2024-08-25"},
            {"Document": "Regulatory Strategy Document", "Status": "Pending Update", "Reviewer": "Mike Wilson", "Due Date": "2024-08-30"},
            {"Document": "Quality Management Procedures", "Status": "Approved", "Reviewer": "Lisa Brown", "Due Date": "2024-08-15"}
        ]
        
        df_doc_status = pd.DataFrame(doc_status_data)
        st.dataframe(df_doc_status, use_container_width=True, hide_index=True)
    
    with tab5:
        st.header("⚡ Workflow Automation")
        
        st.subheader("🔄 Automated Workflows")
        
        # Workflow configuration
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("""
            ### 📋 Change Detection Workflow
            1. **Monitor** regulatory sources (24/7)
            2. **Analyze** change impact using AI
            3. **Map** to affected products/documents
            4. **Trigger** automated actions
            5. **Notify** stakeholders
            6. **Track** compliance status
            """)
            
            st.markdown("""
            ### 📄 Document Update Workflow  
            1. **Detect** regulatory change
            2. **Identify** affected documents
            3. **Generate** tracked changes
            4. **Route** for review/approval
            5. **Update** version control
            6. **Distribute** final versions
            """)
        
        with col2:
            st.markdown("""
            ### 🎯 Compliance Monitoring
            1. **Assess** regulatory requirements
            2. **Compare** current practices
            3. **Identify** gaps
            4. **Generate** action plans
            5. **Monitor** implementation
            6. **Report** compliance status
            """)
            
            st.markdown("""
            ### 🔔 Notification Workflow
            1. **Prioritize** changes by impact
            2. **Route** to responsible teams
            3. **Escalate** high-priority items
            4. **Track** acknowledgments
            5. **Follow-up** on actions
            6. **Archive** completed items
            """)
        
        # Workflow automation settings
        st.subheader("⚙️ Automation Configuration")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown("**Change Detection**")
            scan_frequency = st.selectbox("Scan Frequency", ["Real-time", "Hourly", "Daily", "Weekly"])
            ai_analysis = st.checkbox("AI Impact Analysis", value=True)
            auto_classification = st.checkbox("Auto Classification", value=True)
        
        with col2:
            st.markdown("**Document Management**")
            auto_doc_update = st.checkbox("Auto Document Updates", value=True)
            version_control = st.checkbox("Version Control Integration", value=True)
            approval_routing = st.checkbox("Approval Routing", value=True)
        
        with col3:
            st.markdown("**Notifications**")
            email_alerts = st.checkbox("Email Notifications", value=True)
            slack_integration = st.checkbox("Slack Integration", value=False)
            dashboard_alerts = st.checkbox("Dashboard Alerts", value=True)
        
        # Workflow status dashboard
        st.subheader("📊 Active Workflows")
        
        workflow_data = [
            {"Workflow ID": "WF-001", "Type": "Change Detection", "Status": "Running", "Last Execution": "2024-08-12 14:30", "Next Run": "2024-08-12 15:30"},
            {"Workflow ID": "WF-002", "Type": "Document Update", "Status": "Pending Review", "Last Execution": "2024-08-12 09:15", "Next Run": "Manual"},
            {"Workflow ID": "WF-003", "Type": "Compliance Check", "Status": "Completed", "Last Execution": "2024-08-12 08:00", "Next Run": "2024-08-13 08:00"},
            {"Workflow ID": "WF-004", "Type": "Notification", "Status": "Running", "Last Execution": "2024-08-12 14:45", "Next Run": "Real-time"}
        ]
        
        df_workflows = pd.DataFrame(workflow_data)
        st.dataframe(df_workflows, use_container_width=True, hide_index=True)
        
        # Integration settings
        st.subheader("🔗 System Integrations")
        
        tab_int1, tab_int2, tab_int3 = st.tabs(["Document Management", "Communication", "Project Management"])
        
        with tab_int1:
            st.write("**SharePoint Integration**")
            sharepoint_url = st.text_input("SharePoint Site URL", placeholder="https://company.sharepoint.com/sites/regulatory")
            sharepoint_folder = st.text_input("Document Library", placeholder="Regulatory Documents")
            
            st.write("**Version Control**")
            vc_system = st.selectbox("Version Control System", ["Git", "SharePoint Versions", "Custom"])
            auto_commit = st.checkbox("Auto-commit document changes")
        
        with tab_int2:
            st.write("**Email Configuration**")
            smtp_server = st.text_input("SMTP Server", placeholder="smtp.company.com")
            notification_template = st.selectbox("Email Template", ["Standard", "Urgent", "Custom"])
            
            st.write("**Slack Integration**")
            slack_webhook = st.text_input("Slack Webhook URL", placeholder="https://hooks.slack.com/services/...")
            slack_channel = st.text_input("Default Channel", placeholder="#regulatory-alerts")
        
        with tab_int3:
            st.write("**JIRA Integration**")
            jira_url = st.text_input("JIRA Instance URL", placeholder="https://company.atlassian.net")
            jira_project = st.text_input("Project Key", placeholder="REG")
            auto_create_tickets = st.checkbox("Auto-create tickets for high-impact changes")
            
            st.write("**Task Assignment**")
            default_assignee = st.selectbox("Default Assignee", ["Regulatory Team Lead", "Quality Manager", "Compliance Officer"])
    
    with tab6:
        st.header("🎯 Compliance Monitoring")
        
        # Compliance dashboard
        st.subheader("📊 Compliance Overview")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Compliance Score", "87%", "↑ 3%")
        with col2:
            st.metric("Open Actions", "12", "↓ 5")
        with col3:
            st.metric("Overdue Items", "2", "↓ 1")
        with col4:
            st.metric("Risk Level", "Medium", "↓")
        
        # Compliance by region
        st.subheader("🌍 Regional Compliance Status")
        
        compliance_data = {
            "Region": ["US", "EU", "India", "International"],
            "Compliance Score": [92, 85, 78, 90],
            "Open Actions": [3, 5, 6, 2],
            "Last Assessment": ["2024-08-10", "2024-08-08", "2024-08-05", "2024-08-12"]
        }
        
        df_compliance = pd.DataFrame(compliance_data)
        
        # Compliance score chart
        fig_compliance = px.bar(
            df_compliance,
            x="Region",
            y="Compliance Score",
            color="Compliance Score",
            color_continuous_scale="RdYlGn",
            title="Compliance Scores by Region"
        )
        fig_compliance.update_layout(showlegend=False)
        st.plotly_chart(fig_compliance, use_container_width=True)
        
        # Detailed compliance table
        st.dataframe(df_compliance, use_container_width=True, hide_index=True)
        
        # Gap analysis
        st.subheader("🔍 Compliance Gap Analysis")
        
        gap_data = [
            {"Area": "ICH M10 Implementation", "Current Status": "Partial", "Required By": "2024-12-31", "Gap": "High", "Action": "Update 15 SOPs"},
            {"Area": "EU MDR Compliance", "Current Status": "In Progress", "Required By": "2024-10-15", "Gap": "Medium", "Action": "Complete technical files"},
            {"Area": "India Schedule Y Updates", "Current Status": "Not Started", "Required By": "2024-11-30", "Gap": "High", "Action": "Revise clinical protocols"},
            {"Area": "FDA DSCSA Requirements", "Current Status": "Compliant", "Required By": "N/A", "Gap": "Low", "Action": "Monitor updates"}
        ]
        
        df_gaps = pd.DataFrame(gap_data)
        st.dataframe(df_gaps, use_container_width=True, hide_index=True)
        
        # Action plan generator
        st.subheader("📋 Automated Action Plans")
        
        if st.button("Generate Compliance Action Plans"):
            with st.spinner("Analyzing compliance gaps and generating action plans..."):
                time.sleep(2)
                
                st.success("✅ Action plans generated!")
                
                for gap in gap_data[:2]:  # Show first 2 for demo
                    with st.expander(f"Action Plan: {gap['Area']}"):
                        st.write(f"**Current Status:** {gap['Current Status']}")
                        st.write(f"**Required By:** {gap['Required By']}")
                        st.write(f"**Risk Level:** {gap['Gap']}")
                        
                        st.write("**Recommended Actions:**")
                        if gap['Area'] == "ICH M10 Implementation":
                            st.write("1. Audit current bioanalytical SOPs against ICH M10 requirements")
                            st.write("2. Identify specific gaps in method validation procedures")
                            st.write("3. Update SOPs with new requirements (matrix effect, incurred sample reanalysis)")
                            st.write("4. Train analytical team on updated procedures")
                            st.write("5. Validate updated methods")
                        else:
                            st.write("1. Review current EU MDR documentation")
                            st.write("2. Complete missing technical file sections")
                            st.write("3. Update risk management procedures")
                            st.write("4. Submit updated files to notified body")
                        
                        st.write(f"**Estimated Timeline:** 8-12 weeks")
                        st.write(f"**Resource Requirements:** 2-3 FTEs")
        
        # Risk assessment
        st.subheader("⚠️ Risk Assessment")
        
        risk_data = [
            {"Risk": "Non-compliance with ICH M10", "Probability": "Medium", "Impact": "High", "Risk Score": 8, "Mitigation": "Accelerate SOP updates"},
            {"Risk": "Delayed EU MDR submission", "Probability": "Low", "Impact": "High", "Risk Score": 6, "Mitigation": "Dedicated project team"},
            {"Risk": "India regulatory delays", "Probability": "High", "Impact": "Medium", "Risk Score": 7, "Mitigation": "Engage local consultant"},
            {"Risk": "FDA inspection findings", "Probability": "Low", "Impact": "High", "Risk Score": 5, "Mitigation": "Pre-inspection audit"}
        ]
        
        df_risks = pd.DataFrame(risk_data)
        
        # Risk matrix visualization
        fig_risk = px.scatter(
            df_risks,
            x=["Low", "Low", "High", "Low"],  # Probability mapping
            y=["High", "High", "Medium", "High"],  # Impact mapping
            size="Risk Score",
            color="Risk Score",
            hover_name="Risk",
            title="Risk Assessment Matrix"
        )
        st.plotly_chart(fig_risk, use_container_width=True)
        
        st.dataframe(df_risks, use_container_width=True, hide_index=True)
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666;">
        <p>Pharma Regulatory Intelligence Platform v2.0 | 
        Last Updated: August 12, 2024 | 
        <a href="mailto:regulatory@company.com">Contact Support</a></p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
